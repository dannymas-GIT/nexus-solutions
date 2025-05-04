import os
import json
import logging
import mimetypes
from pathlib import Path
from typing import Dict, Any, List, Optional
from datetime import datetime
from io import BytesIO

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from ...core.config import settings
from .base import DocumentService
from ..storage import get_storage_service

logger = logging.getLogger(__name__)

class DocxDocumentService(DocumentService):
    """Service for generating Word documents."""
    
    def __init__(self):
        self.templates_path = Path(settings.DOCX_TEMPLATES_PATH)
        self.storage_service = get_storage_service()
        
        # Ensure the templates directory exists
        self._ensure_templates_path_exists()
    
    def _ensure_templates_path_exists(self):
        """Ensure that the templates directory exists."""
        if not self.templates_path.exists():
            logger.info(f"Creating templates directory: {self.templates_path}")
            self.templates_path.mkdir(parents=True, exist_ok=True)
    
    def _get_template_path(self, template_name: str) -> Path:
        """Get the full path to a template."""
        template_path = self.templates_path / f"{template_name}.docx"
        if not template_path.exists():
            raise FileNotFoundError(f"Template not found: {template_path}")
        return template_path
    
    def _get_schema_path(self, template_name: str) -> Path:
        """Get the path to a template's schema file."""
        schema_path = self.templates_path / f"{template_name}.schema.json"
        return schema_path
    
    def _get_info_path(self, template_name: str) -> Path:
        """Get the path to a template's info file."""
        info_path = self.templates_path / f"{template_name}.info.json"
        return info_path
    
    async def generate_document(
        self, 
        template_name: str, 
        data: Dict[str, Any], 
        output_path: Optional[str] = None
    ) -> Path:
        """
        Generate a Word document from a template and data.
        
        Args:
            template_name: Name of the template to use
            data: Data to populate the template with
            output_path: Optional path where to save the document
            
        Returns:
            Path to the generated document
        """
        template_path = self._get_template_path(template_name)
        
        # Create a document from the template
        doc = Document(template_path)
        
        # Process paragraphs
        for paragraph in doc.paragraphs:
            self._process_paragraph(paragraph, data)
        
        # Process tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._process_paragraph(paragraph, data)
        
        # Create output filename if not provided
        if not output_path:
            timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
            output_path = f"documents/{template_name}_{timestamp}.docx"
        
        # Ensure directories exist
        output_file = Path(output_path)
        output_file.parent.mkdir(parents=True, exist_ok=True)
        
        # Save the document
        doc.save(output_file)
        logger.info(f"Generated document: {output_file}")
        
        return output_file
    
    def _process_paragraph(self, paragraph, data: Dict[str, Any]):
        """
        Process a paragraph in a document, replacing placeholders with data.
        
        Args:
            paragraph: The paragraph to process
            data: The data dictionary
        """
        if not paragraph.runs:
            return
        
        # Get the paragraph text
        text = paragraph.text
        
        # Check if there are any placeholders in the paragraph
        if "{{" not in text or "}}" not in text:
            return
        
        # Replace placeholders
        new_text = text
        for key, value in data.items():
            placeholder = f"{{{{{key}}}}}"
            if placeholder in new_text:
                new_text = new_text.replace(placeholder, str(value))
        
        # Remove all runs except the first one
        for _ in range(len(paragraph.runs) - 1):
            p = paragraph._p
            p.remove(paragraph.runs[1]._r)
        
        # Set the new text to the first run
        paragraph.runs[0].text = new_text
    
    async def list_templates(self) -> List[Dict[str, Any]]:
        """
        List available Word templates.
        
        Returns:
            List of template metadata
        """
        templates = []
        
        if not self.templates_path.exists():
            return templates
        
        for file_path in self.templates_path.glob("*.docx"):
            if file_path.is_file() and not file_path.name.startswith("."):
                template_name = file_path.stem
                
                # Try to get info from the info file
                info = await self.get_template_info(template_name)
                
                templates.append({
                    "name": template_name,
                    "file": file_path.name,
                    "type": "docx",
                    "info": info
                })
        
        return templates
    
    async def get_template_info(self, template_name: str) -> Dict[str, Any]:
        """
        Get information about a specific template.
        
        Args:
            template_name: Name of the template
            
        Returns:
            Template metadata
        """
        info_path = self._get_info_path(template_name)
        
        # Default info
        info = {
            "name": template_name,
            "type": "docx",
            "description": f"Word template: {template_name}",
            "pages": 0
        }
        
        # Try to get info from the info file
        if info_path.exists():
            try:
                with open(info_path, "r") as f:
                    file_info = json.load(f)
                    info.update(file_info)
            except Exception as e:
                logger.error(f"Error reading template info file: {e}")
        
        return info
    
    async def get_template_schema(self, template_name: str) -> Dict[str, Any]:
        """
        Get the schema of data required for a specific template.
        
        Args:
            template_name: Name of the template
            
        Returns:
            JSON schema of the data required for the template
        """
        schema_path = self._get_schema_path(template_name)
        
        # Default schema
        schema = {
            "type": "object",
            "properties": {},
            "required": []
        }
        
        # Try to get schema from the schema file
        if schema_path.exists():
            try:
                with open(schema_path, "r") as f:
                    file_schema = json.load(f)
                    schema.update(file_schema)
            except Exception as e:
                logger.error(f"Error reading template schema file: {e}")
        else:
            # If no schema file exists, try to extract placeholders from the template
            try:
                template_path = self._get_template_path(template_name)
                doc = Document(template_path)
                
                # Find all placeholders in paragraphs and tables
                placeholders = set()
                
                # Check paragraphs
                for paragraph in doc.paragraphs:
                    text = paragraph.text
                    import re
                    for match in re.finditer(r"{{([^}]+)}}", text):
                        placeholders.add(match.group(1))
                
                # Check tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                text = paragraph.text
                                for match in re.finditer(r"{{([^}]+)}}", text):
                                    placeholders.add(match.group(1))
                
                # Create schema properties for each placeholder
                for placeholder in placeholders:
                    schema["properties"][placeholder] = {"type": "string"}
                    schema["required"].append(placeholder)
            
            except Exception as e:
                logger.error(f"Error extracting placeholders from template: {e}")
        
        return schema 